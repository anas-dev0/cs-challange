"""
Resume/CV Parser Module - FIXED VERSION
Extracts and cleans text from PDF, DOCX, and TXT files.
"""

import re
import os
from typing import Optional
from io import BytesIO

# Import libraries for different file formats
try:
    import PyPDF2  # For PDF files
except ImportError:
    PyPDF2 = None

try:
    from docx import Document  # For DOCX files
except ImportError:
    Document = None

try:
    import fitz  # PyMuPDF - better PDF extraction
except ImportError:
    fitz = None

try:
    from PIL import Image
except ImportError:
    Image = None


def rectangles_overlap(rect1, rect2, tolerance=5):
    """
    Check if two rectangles overlap with some tolerance.
    
    Args:
        rect1: (x0, y0, x1, y1) - first rectangle
        rect2: (x0, y0, x1, y1) or Rect object - second rectangle
        tolerance: pixels of tolerance for overlap detection
    
    Returns:
        bool: True if rectangles overlap
    """
    # Handle fitz.Rect objects
    if hasattr(rect2, 'x0'):
        r2 = (rect2.x0, rect2.y0, rect2.x1, rect2.y1)
    else:
        r2 = rect2
    
    # Expand rectangles slightly for tolerance
    r1_expanded = (rect1[0] - tolerance, rect1[1] - tolerance, 
                   rect1[2] + tolerance, rect1[3] + tolerance)
    
    # Check if rectangles overlap
    return not (r1_expanded[2] < r2[0] or  # rect1 is left of rect2
                r1_expanded[0] > r2[2] or  # rect1 is right of rect2
                r1_expanded[3] < r2[1] or  # rect1 is above rect2
                r1_expanded[1] > r2[3])    # rect1 is below rect2


def parse_document_with_images(file_path: str) -> dict:
    """
    Parse a resume/CV file and extract clean text content along with page images.
    
    Args:
        file_path (str): Path to the document file (.pdf, .docx, or .txt)
    
    Returns:
        dict: Dictionary containing 'text' and 'images' keys
    """
    result = {
        'text': None,
        'images': []
    }
    
    # First get the text
    text = parse_document(file_path)
    result['text'] = text
    
    # If it's a PDF, also get images
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    
    if file_extension == '.pdf':
        print("🖼️  Extracting PDF pages as images for visual analysis...")
        images = pdf_to_images(file_path)
        result['images'] = images
    
    return result


def parse_document(file_path: str) -> Optional[str]:
    """
    Parse a resume/CV file and extract clean text content.
    
    Args:
        file_path (str): Path to the document file (.pdf, .docx, or .txt)
    
    Returns:
        str: Cleaned text content from the document
        None: If parsing fails or file type is not supported
    """
    result = parse_document_with_metadata(file_path)
    return result['text'] if result else None

def parse_document_with_metadata(file_path: str) -> Optional[dict]:
    """
    Parse a resume/CV file and extract clean text content with metadata (links).
    
    Args:
        file_path (str): Path to the document file (.pdf, .docx, or .txt)
    
    Returns:
        dict: Dictionary with 'text' and 'links' keys
        None: If parsing fails or file type is not supported
    """
    
    print(f"\n🔍 DEBUG: Starting to parse: {file_path}")
    
    # Check if file exists
    if not os.path.exists(file_path):
        print(f"❌ Error: File not found - {file_path}")
        return None
    
    print(f"✅ File exists")
    
    # Get file extension to determine file type
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    print(f"📄 File extension: {file_extension}")
    
    # Check file size
    file_size = os.path.getsize(file_path)
    print(f"📏 File size: {file_size} bytes")
    
    # Extract raw text based on file type
    try:
        result = {'text': None, 'links': []}
        
        if file_extension == '.pdf':
            print("🔄 Attempting PDF extraction...")
            pdf_result = extract_text_and_links_from_pdf(file_path)
            result['text'] = pdf_result['text']
            result['links'] = pdf_result['links']
        elif file_extension in ['.docx', '.doc']:
            print("🔄 Attempting DOCX extraction...")
            result['text'] = extract_text_from_docx(file_path)
            result['links'] = []  # DOCX link extraction can be added later if needed
        elif file_extension == '.txt':
            print("🔄 Attempting TXT extraction...")
            result['text'] = extract_text_from_txt(file_path)
            result['links'] = []
        else:
            print(f"❌ Error: Unsupported file type - {file_extension}")
            return None
        
        # Check if text extraction was successful
        if result['text'] is None:
            print("❌ Error: Failed to extract text from document")
            return None
        
        print(f"✅ Extracted {len(result['text'])} characters")
        print(f"🔗 Found {len(result['links'])} links")
        print(f"📝 First 100 chars: {result['text'][:100]}")
        
        # Clean the extracted text
        result['text'] = clean_text(result['text'])
        print(f"✅ Cleaned text: {len(result['text'])} characters")
        return result
    
    except Exception as e:
        print(f"❌ Error parsing document: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from PDF - tries PyMuPDF first (better), falls back to PyPDF2
    
    Args:
        file_path (str): Path to the PDF file
    
    Returns:
        str: Raw text extracted from PDF
        None: If extraction fails
    """
    return extract_text_and_links_from_pdf(file_path)['text']


def extract_text_and_links_from_pdf(file_path: str) -> dict:
    """
    Extract text and links from PDF using PyMuPDF.
    Embeds link URLs directly into the text where they appear for better context.
    
    Args:
        file_path (str): Path to the PDF file
    
    Returns:
        dict: Dictionary with 'text' (with inline link annotations) and 'links' keys
    """
    result = {
        'text': "",
        'links': []
    }
    
    # Try PyMuPDF first (much more reliable for complex PDFs)
    if fitz is not None:
        try:
            print("   Using PyMuPDF for extraction with inline link annotations...")
            
            with fitz.open(file_path) as doc:
                print(f"   PDF has {len(doc)} pages")
                for page_num, page in enumerate(doc):
                    # Extract links first
                    page_links = []
                    try:
                        links = page.get_links()
                        for link in links:
                            if 'uri' in link and link['uri']:
                                link_rect = link.get('from')
                                page_links.append({
                                    'uri': link['uri'],
                                    'rect': link_rect
                                })
                                if link['uri'] not in result['links']:
                                    result['links'].append(link['uri'])
                        print(f"   Page {page_num + 1}: found {len(page_links)} links")
                    except Exception as link_error:
                        print(f"   ⚠️  Could not extract links from page {page_num + 1}: {str(link_error)}")
                    
                    # Extract text with position information
                    try:
                        blocks = page.get_text("dict")["blocks"]
                        page_text = ""
                        
                        for block in blocks:
                            if block.get("type") == 0:  # Text block
                                for line in block.get("lines", []):
                                    line_text = ""
                                    for span in line.get("spans", []):
                                        span_text = span.get("text", "")
                                        span_bbox = span.get("bbox", (0, 0, 0, 0))
                                        
                                        # Check if this span overlaps with any link
                                        matching_link = None
                                        for link_info in page_links:
                                            if rectangles_overlap(span_bbox, link_info['rect']):
                                                matching_link = link_info['uri']
                                                break
                                        
                                        # Add text with link annotation if found
                                        if matching_link and span_text.strip():
                                            # Embed the URL right after the linked text
                                            line_text += f"{span_text} <{matching_link}> "
                                        else:
                                            line_text += span_text
                                    
                                    page_text += line_text.strip() + "\n"
                        
                        result['text'] += page_text + "\n"
                        print(f"   Page {page_num + 1}: extracted {len(page_text)} chars with {len(page_links)} inline link annotations")
                    
                    except Exception as text_error:
                        # Fallback to simple text extraction if detailed parsing fails
                        print(f"   ⚠️  Detailed text extraction failed, using simple method: {str(text_error)}")
                        page_text = page.get_text()
                        result['text'] += page_text + "\n"
                        print(f"   Page {page_num + 1}: extracted {len(page_text)} chars (simple mode)")
            
            if result['text'].strip():
                print(f"   ✅ PyMuPDF successfully extracted {len(result['text'])} chars with {len(result['links'])} links annotated inline")
                return result
            else:
                print("   ⚠️  PyMuPDF extracted empty text, trying PyPDF2...")
        
        except Exception as e:
            print(f"   ⚠️  PyMuPDF error: {str(e)}, trying PyPDF2...")
            import traceback
            traceback.print_exc()
    else:
        print("   ⚠️  PyMuPDF not available, using PyPDF2...")
    
    # Fallback to PyPDF2
    if PyPDF2 is None:
        print("   ❌ PyPDF2 not installed. Install with: pip install PyPDF2")
        return result if result['text'] else {'text': None, 'links': []}
    
    try:
        print("   Using PyPDF2 for extraction...")
        text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            print(f"   PDF has {len(pdf_reader.pages)} pages")
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text() or ""
                print(f"   Page {page_num + 1}: extracted {len(page_text)} chars")
                text += page_text + "\n"
        
        if text.strip():
            print(f"   ✅ PyPDF2 successfully extracted {len(text)} chars")
            result['text'] = text
            # Note: PyPDF2 doesn't easily extract links, so we keep links empty
            print(f"   ⚠️  PyPDF2 doesn't support link extraction - use PyMuPDF for full features")
            return result
        else:
            print("   ⚠️  PyPDF2 extracted empty text - PDF might be image-based (scanned)")
            return {'text': None, 'links': []}
    
    except Exception as e:
        print(f"   ❌ PyPDF2 error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {'text': None, 'links': []}


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from DOCX file with better error handling
    
    Args:
        file_path (str): Path to the DOCX file
    
    Returns:
        str: Raw text extracted from DOCX
        None: If extraction fails
    """
    if Document is None:
        print("   ❌ python-docx not installed. Install with: pip install python-docx")
        return None
    
    try:
        print("   Using python-docx for extraction...")
        doc = Document(file_path)
        
        text = ""
        
        # Extract text from paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        print(f"   ✅ Extracted {len(text)} characters from DOCX")
        
        if not text.strip():
            print("   ⚠️  DOCX appears to be empty")
            return None
            
        return text
    
    except Exception as e:
        print(f"   ❌ DOCX extraction error: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def extract_text_from_txt(file_path: str) -> Optional[str]:
    """
    Extract text from a TXT file.
    
    Args:
        file_path (str): Path to the TXT file
    
    Returns:
        str: Raw text from TXT file
        None: If extraction fails
    """
    
    try:
        # Try UTF-8 first
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        print(f"   ✅ Read TXT file with UTF-8 encoding")
        return text
    
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            print(f"   ✅ Read TXT file with latin-1 encoding")
            return text
        except Exception as e:
            print(f"   ❌ Error reading TXT file: {str(e)}")
            return None
    
    except Exception as e:
        print(f"   ❌ Error reading TXT file: {str(e)}")
        return None


def pdf_to_images(file_path: str, output_dir: str = None) -> list:
    """
    Convert PDF pages to images for vision-based analysis.
    
    Args:
        file_path (str): Path to the PDF file
        output_dir (str): Directory to save images (optional)
    
    Returns:
        list: List of PIL Image objects or image paths
    """
    if not fitz:
        print("   ⚠️  PyMuPDF not available for PDF to image conversion")
        return []
    
    try:
        print(f"   Converting PDF to images: {file_path}")
        images = []
        
        with fitz.open(file_path) as doc:
            for page_num in range(len(doc)):
                page = doc[page_num]
                # Render page to pixmap (image)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
                
                # Convert to PIL Image
                img_data = pix.tobytes("png")
                img = Image.open(BytesIO(img_data)) if Image else None
                
                if img:
                    if output_dir:
                        img_path = os.path.join(output_dir, f"page_{page_num + 1}.png")
                        img.save(img_path)
                        images.append(img_path)
                        print(f"   Saved page {page_num + 1} to {img_path}")
                    else:
                        images.append(img)
                        print(f"   Converted page {page_num + 1} to image")
        
        print(f"   ✅ Successfully converted {len(images)} pages to images")
        return images
    
    except Exception as e:
        print(f"   ❌ Error converting PDF to images: {str(e)}")
        import traceback
        traceback.print_exc()
        return []


def clean_text(text: str) -> str:
    """
    Clean extracted text - gentler approach that preserves more content
    
    Args:
        text (str): Raw text to clean
    
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Replace multiple spaces/tabs with single space (preserve newlines)
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Replace 3+ newlines with just 2 (preserve paragraph breaks)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove leading and trailing whitespace from entire text
    text = text.strip()
    
    return text


# Test script
if __name__ == "__main__":
    import sys
    
    print("="*60)
    print("CV PARSER TEST")
    print("="*60)
    
    # Check installed libraries
    print("\n📦 Checking libraries:")
    if PyPDF2:
        print("   ✅ PyPDF2 installed")
    else:
        print("   ❌ PyPDF2 NOT installed - pip install PyPDF2")
    
    if Document:
        print("   ✅ python-docx installed")
    else:
        print("   ❌ python-docx NOT installed - pip install python-docx")
    
    if fitz:
        print("   ✅ PyMuPDF installed (recommended)")
    else:
        print("   ⚠️  PyMuPDF NOT installed (optional) - pip install PyMuPDF")
    
    # Test with file
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
    else:
        test_file = input("\n📁 Enter file path to test: ").strip()
    
    if not os.path.exists(test_file):
        print(f"\n❌ File not found: {test_file}")
        sys.exit(1)
    
    print(f"\n{'='*60}")
    result = parse_document(test_file)
    print(f"{'='*60}")
    
    if result:
        print(f"\n✅ SUCCESS!")
        print(f"📏 Total length: {len(result)} characters")
        print(f"📏 Word count: {len(result.split())}")
        print(f"\n📝 First 300 characters:\n")
        print(result[:300])
        print("\n...")
        if len(result) > 500:
            print(f"\n📝 Last 200 characters:\n")
            print(result[-200:])
    else:
        print("\n❌ FAILED to parse document")
        print("\nTroubleshooting:")
        print("1. Make sure libraries are installed:")
        print("   pip install PyPDF2 python-docx PyMuPDF")
        print("2. Check if the file is corrupted")
        print("3. For scanned PDFs, you need OCR (tesseract)")