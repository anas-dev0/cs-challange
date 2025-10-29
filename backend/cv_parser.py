"""
Fast and reliable CV parser for PDF and DOCX files.
Uses PyMuPDF (fitz) for PDF extraction - no external dependencies needed.
"""
import os
import io
from typing import Optional

# Import PDF library
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    print("ERROR: PyMuPDF not installed. Run: pip install PyMuPDF")

# Import DOCX library
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Import OCR library (optional)
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


def extract_text_from_cv(file_path: str) -> str:
    """
    Fast extraction of text from CV files (PDF or DOCX).
    
    Args:
        file_path: Path to the CV file
        
    Returns:
        Extracted text content, or empty string if extraction fails
    """
    if not os.path.exists(file_path):
        print(f"ERROR: File not found: {file_path}")
        return ""
    
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    print(f"📄 Parsing {ext} file: {os.path.basename(file_path)}")
    
    try:
        if ext == ".pdf":
            return _extract_pdf_fast(file_path)
        elif ext == ".docx":
            return _extract_docx(file_path)
        else:
            print(f"ERROR: Unsupported file type: {ext}")
            return ""
    except Exception as e:
        print(f"ERROR: Failed to extract text: {str(e)}")
        import traceback
        traceback.print_exc()
        return ""


def _extract_pdf_fast(file_path: str) -> str:
    """
    Fast PDF text extraction using PyMuPDF.
    Tries multiple extraction methods for compatibility.
    """
    if not HAS_PYMUPDF:
        print("ERROR: PyMuPDF not available")
        return ""
    
    doc = None
    
    try:
        doc = fitz.open(file_path)
        page_count = len(doc)
        print(f"📖 PDF has {page_count} pages")
        
        # Method 1: Try standard text extraction
        text_parts = []
        for page_num in range(page_count):
            page = doc[page_num]
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text.strip())
        
        if text_parts:
            result = "\n\n".join(text_parts)
            print(f"✅ Extracted {len(result)} characters from {page_count} pages (method: text)")
            return result
        
        # Method 2: Try blocks extraction (better for complex layouts)
        print("⚠️ No text found, trying blocks method...")
        text_parts = []
        for page_num in range(page_count):
            page = doc[page_num]
            blocks = page.get_text("blocks")
            for block in blocks:
                if len(block) >= 5:  # block format: (x0, y0, x1, y1, "text", block_no, block_type)
                    block_text = block[4]
                    if block_text.strip():
                        text_parts.append(block_text.strip())
        
        if text_parts:
            result = "\n".join(text_parts)
            print(f"✅ Extracted {len(result)} characters from {page_count} pages (method: blocks)")
            return result
        
        # Method 3: Try dict extraction (most detailed)
        print("⚠️ No text in blocks, trying dict method...")
        text_parts = []
        for page_num in range(page_count):
            page = doc[page_num]
            page_dict = page.get_text("dict")
            page_text = _extract_text_from_dict(page_dict)
            if page_text.strip():
                text_parts.append(page_text.strip())
        
        if text_parts:
            result = "\n\n".join(text_parts)
            print(f"✅ Extracted {len(result)} characters from {page_count} pages (method: dict)")
            return result
        
        # Method 4: Try words extraction
        print("⚠️ No text in dict, trying words method...")
        text_parts = []
        for page_num in range(page_count):
            page = doc[page_num]
            words = page.get_text("words")
            if words:
                page_text = " ".join([w[4] for w in words if len(w) >= 5])
                if page_text.strip():
                    text_parts.append(page_text.strip())
        
        if text_parts:
            result = "\n\n".join(text_parts)
            print(f"✅ Extracted {len(result)} characters from {page_count} pages (method: words)")
            return result
        
        # Method 5: Try OCR for image-based PDFs
        if HAS_OCR:
            print("⚠️ No text found, trying OCR extraction...")
            return _extract_pdf_with_ocr(doc, page_count)
        
        print(f"⚠️ Could not extract text from PDF - may be image-based or scanned")
        print(f"💡 Install OCR support: pip install pytesseract pillow")
        print(f"💡 Also install Tesseract: https://github.com/tesseract-ocr/tesseract")
        return ""
        
    except Exception as e:
        print(f"ERROR: PDF extraction failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return ""
    finally:
        if doc is not None:
            try:
                doc.close()
            except:
                pass


def _extract_text_from_dict(page_dict: dict) -> str:
    """
    Extract text from PyMuPDF's dict format.
    Preserves better structure than raw text extraction.
    """
    text_parts = []
    
    try:
        for block in page_dict.get("blocks", []):
            if block.get("type") == 0:  # Text block
                for line in block.get("lines", []):
                    line_text = ""
                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                    if line_text.strip():
                        text_parts.append(line_text.strip())
        
        return "\n".join(text_parts)
    except:
        return ""


def _extract_docx(file_path: str) -> str:
    """
    Extract text from DOCX files using python-docx.
    """
    if not HAS_DOCX:
        print("ERROR: python-docx not available")
        return ""
    
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        result = "\n".join(text_parts)
        print(f"✅ Extracted {len(result)} characters")
        return result
        
    except Exception as e:
        print(f"ERROR: DOCX extraction failed: {str(e)}")
        return ""


def _extract_pdf_with_ocr(doc, page_count: int) -> str:
    """
    Extract text from image-based PDF using OCR.
    This is slower but works for scanned PDFs.
    """
    if not HAS_OCR:
        return ""
    
    try:
        # Set Tesseract path for Windows
        if os.name == 'nt':
            tesseract_paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            ]
            for path in tesseract_paths:
                if os.path.exists(path):
                    pytesseract.pytesseract.tesseract_cmd = path
                    break
        
        text_parts = []
        # Limit to first 3 pages for speed
        max_pages = min(3, page_count)
        
        for page_num in range(max_pages):
            page = doc[page_num]
            
            # Render page as image at 2x resolution for better OCR
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            
            # Convert to PIL Image
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            
            # Perform OCR
            page_text = pytesseract.image_to_string(img, lang='eng')
            
            if page_text.strip():
                text_parts.append(page_text.strip())
                print(f"  ✅ Page {page_num + 1}: {len(page_text)} chars (OCR)")
        
        if text_parts:
            result = "\n\n".join(text_parts)
            print(f"✅ Extracted {len(result)} characters using OCR from {max_pages} pages")
            return result
        
        return ""
        
    except Exception as e:
        print(f"ERROR: OCR extraction failed: {str(e)}")
        return ""
